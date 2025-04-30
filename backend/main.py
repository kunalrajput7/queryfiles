from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import firebase_admin
from sentence_transformers import SentenceTransformer
from google.cloud import storage, firestore
from google.cloud import vision_v1
import faiss
import fitz  # PyMuPDF
import numpy as np
import requests
import os
import uuid
import json
import time
from docx import Document
import openpyxl
import xlrd
from firebase_admin import auth
from firebase_admin import credentials
from pydantic import BaseModel


class UIDRequest(BaseModel):
    uid: str

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "https://<your-vercel-domain>.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "../gcs_key.json"
BUCKET_NAME = "pdf-faiss-bucket"
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-caa811add2ff4ee683ec33aba69fa270")  # Set in environment
cred = credentials.Certificate("../gcs_key.json")
firebase_admin.initialize_app(cred, {
    "storageBucket": "pdf-faiss-bucket"
})

storage_client = storage.Client()
bucket = storage_client.bucket(BUCKET_NAME)
firestore_client = firestore.Client()
vision_client = vision_v1.ImageAnnotatorClient()

# Use CPU for embedding (no GPU needed)
device = "cpu"
print(f"Using device: {device}")

embedder = SentenceTransformer("all-MiniLM-L6-v2", device=device)
print("Warming up embedder...")
embedder.encode(["Warm-up query"], convert_to_numpy=True)
print("Embedder warmed up.")

# Chat history for context-awareness
chat_history = {}
active_indices = {}
active_chunks_dict = {}

def generate_uid():
    return str(uuid.uuid4())

def upload_to_gcs(source_file, destination_blob_name):
    try:
        blob = bucket.blob(destination_blob_name)
        blob.upload_from_filename(source_file)
        print(f"File {source_file} uploaded to gs://{BUCKET_NAME}/{destination_blob_name}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload to GCS: {str(e)}")

def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text = ""
        total_chars = 0
        page_count = len(doc)
        
        # Extract text with PyMuPDF
        for page in doc:
            page_text = page.get_text("text")
            text += page_text
            total_chars += len(page_text.strip())
        
        doc.close()
        
        # If little text extracted (<100 chars/page), assume image-based PDF and use Cloud Vision API
        if page_count > 0 and total_chars / page_count < 100:
            print(f"Low text content detected ({total_chars} chars in {page_count} pages). Using Cloud Vision API for OCR.")
            # Upload PDF to GCS temporarily for Vision API
            temp_gcs_path = f"temp/{os.path.basename(file_path)}"
            upload_to_gcs(file_path, temp_gcs_path)
            
            # Call Vision API
            gcs_source = vision_v1.GcsSource(uri=f"gs://{BUCKET_NAME}/{temp_gcs_path}")
            input_config = vision_v1.InputConfig(gcs_source=gcs_source, mime_type="application/pdf")
            feature = vision_v1.Feature(type_=vision_v1.Feature.Type.DOCUMENT_TEXT_DETECTION)
            request = vision_v1.AsyncAnnotateFileRequest(
                features=[feature],
                input_config=input_config,
                output_config=vision_v1.OutputConfig(
                    gcs_destination=vision_v1.GcsDestination(uri=f"gs://{BUCKET_NAME}/temp/output/"),
                    batch_size=10  # Process up to 10 pages per output file
                )
            )
            
            operation = vision_client.async_batch_annotate_files(requests=[request])
            operation.result(timeout=600)  # Wait up to 10 minutes for larger PDFs
            
            # Retrieve all output files from GCS
            ocr_text = ""
            output_blobs = list(bucket.list_blobs(prefix="temp/output/"))
            print(f"Found {len(output_blobs)} output files for OCR processing.")
            for output_blob in output_blobs:
                output_json = output_blob.download_as_string().decode("utf-8")
                responses = json.loads(output_json)
                for response in responses.get("responses", []):
                    if "fullTextAnnotation" in response:
                        ocr_text += response["fullTextAnnotation"]["text"] + "\n"
                # Clean up each output blob
                bucket.delete_blob(output_blob.name)
            
            # Clean up temporary input file
            bucket.delete_blob(temp_gcs_path)
            
            text = ocr_text if ocr_text.strip() else text  # Fallback to PyMuPDF text if OCR fails
        
        if not text.strip():
            raise ValueError("No extractable text found in the PDF.")
        print(f"Extracted {len(text)} characters from {page_count} pages.")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting PDF: {str(e)}")

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        if not text.strip():
            raise ValueError("No extractable text found in the Word document.")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting Word document: {str(e)}")

def extract_text_from_excel(file_path, extension):
    try:
        text = ""
        if extension == ".xlsx":
            wb = openpyxl.load_workbook(file_path, read_only=True)
            for sheet in wb:
                for row in sheet.rows:
                    for cell in row:
                        if cell.value:
                            text += str(cell.value) + "\t"
                    text += "\n"
            wb.close()
        elif extension == ".xls":
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    for col_idx in range(sheet.ncols):
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value:
                            text += str(cell_value) + "\t"
                    text += "\n"
        
        if not text.strip():
            raise ValueError("No extractable text found in the Excel file.")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting Excel file: {str(e)}")

async def save_pdf_locally(upload_file: UploadFile):
    temp_path = f"./{upload_file.filename}"
    with open(temp_path, "wb") as f:
        while chunk := await upload_file.read(4096):
            f.write(chunk)
    return temp_path

def process_pdf_and_store_index(text, uid, pdf_name):
    try:
        words = text.split()
        pdf_chunks = [" ".join(words[i:i + 200]) for i in range(0, len(words), 200)]
        if not pdf_chunks:
            raise Exception("No text extracted from file")

        embeddings = embedder.encode(pdf_chunks, convert_to_numpy=True)
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)

        index_file = f"{pdf_name}.index"
        faiss.write_index(index, index_file)
        index_gcs_path = f"{uid}/index/{index_file}"
        upload_to_gcs(index_file, index_gcs_path)
        os.remove(index_file)
        print(f"FAISS index uploaded and deleted locally: {index_gcs_path}")

        chunks_file = f"{pdf_name}.chunks.json"
        with open(chunks_file, "w") as f:
            json.dump(pdf_chunks, f)
        chunks_gcs_path = f"{uid}/chunks/{chunks_file}"
        upload_to_gcs(chunks_file, chunks_gcs_path)
        os.remove(chunks_file)
        print(f"Chunks uploaded and deleted locally: {chunks_gcs_path}")

        return index_gcs_path, chunks_gcs_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

def save_file_metadata(uid, pdf_name, pdf_gcs_path, index_gcs_path, chunks_gcs_path):
    try:
        fileid = str(uuid.uuid4())
        file_data = {
            "filename": pdf_name,
            "pdfUrl": f"https://storage.googleapis.com/{BUCKET_NAME}/{pdf_gcs_path}",
            "indexUrl": f"https://storage.googleapis.com/{BUCKET_NAME}/{index_gcs_path}",
            "chunksUrl": f"https://storage.googleapis.com/{BUCKET_NAME}/{chunks_gcs_path}",
            "upload_date": firestore.SERVER_TIMESTAMP,
        }
        doc_ref = firestore_client.collection("users").document(uid).collection("files").document(fileid)
        doc_ref.set(file_data)
        print(f"File metadata saved for {uid} under file id {fileid}")
        return fileid
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving metadata to Firestore: {str(e)}")

@app.post("/upload_pdf")
async def upload_pdf(file: UploadFile = File(...), uid: str = Form(None)):
    try:
        if not uid:
            uid = generate_uid()
        pdf_name = file.filename
        extension = os.path.splitext(pdf_name)[1].lower()

        # Validate file extension
        if extension not in [".pdf", ".docx", ".xlsx", ".xls"]:
            raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, Word (.docx), and Excel (.xlsx, .xls) are supported.")

        file_local_path = await save_pdf_locally(file)
        pdf_gcs_path = f"{uid}/pdf/{pdf_name}"
        upload_to_gcs(file_local_path, pdf_gcs_path)

        # Extract text based on file type
        if extension == ".pdf":
            text = extract_text_from_pdf(file_local_path)
        elif extension == ".docx":
            text = extract_text_from_docx(file_local_path)
        elif extension in [".xlsx", ".xls"]:
            text = extract_text_from_excel(file_local_path, extension)

        index_gcs_path, chunks_gcs_path = process_pdf_and_store_index(text, uid, pdf_name)
        fileid = save_file_metadata(uid, pdf_name, pdf_gcs_path, index_gcs_path, chunks_gcs_path)

        os.remove(file_local_path)
        return {
            "message": "File uploaded and processed successfully",
            "uid": uid,
            "id": fileid,
            "filename": pdf_name,
        }
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/load_index")
async def load_index(uid: str = Form(...), fileid: str = Form(...)):
    try:
        doc_ref = firestore_client.collection("users").document(uid).collection("files").document(fileid)
        doc = doc_ref.get()
        if not doc.exists:
            raise HTTPException(status_code=404, detail="File metadata not found")
        data = doc.to_dict()
        index_url = data.get("indexUrl")
        chunks_url = data.get("chunksUrl")

        prefix = f"https://storage.googleapis.com/{BUCKET_NAME}/"
        if not (index_url.startswith(prefix) and chunks_url.startswith(prefix)):
            raise HTTPException(status_code=400, detail="Invalid URL format")
        index_path = index_url[len(prefix):]
        chunks_path = chunks_url[len(prefix):]

        local_index_file = f"./temp_index_{uid}.index"
        blob = bucket.blob(index_path)
        blob.download_to_filename(local_index_file)
        active_indices[uid] = faiss.read_index(local_index_file)
        os.remove(local_index_file)
        print(f"Index for file {fileid} loaded for user {uid}")

        local_chunks_file = f"./temp_chunks_{uid}.json"
        blob = bucket.blob(chunks_path)
        blob.download_to_filename(local_chunks_file)
        with open(local_chunks_file, "r") as f:
            active_chunks_dict[uid] = json.load(f)
        os.remove(local_chunks_file)
        print(f"Chunks for file {fileid} loaded for user {uid}")

        return {"message": "Index and chunks loaded successfully"}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading index and chunks: {str(e)}")

@app.post("/query")
async def query_pdf(query: str = Form(...), uid: str = Form(...)):
    try:
        if uid not in active_indices or active_indices[uid] is None:
            raise HTTPException(status_code=400, detail="No FAISS index loaded. Please select a file first.")
        if uid not in active_chunks_dict or not active_chunks_dict[uid]:
            raise HTTPException(status_code=400, detail="No chunks loaded. Please select a file first.")

        index = active_indices[uid]
        chunks = active_chunks_dict[uid]

        start_time = time.time()

        t1 = time.time()
        query_embedding = embedder.encode([query], convert_to_numpy=True)
        print(f"Embedding time: {time.time() - t1:.2f}s")

        t2 = time.time()
        distances, indices = index.search(query_embedding, k=3)
        print(f"FAISS search time: {time.time() - t2:.2f}s")

        t3 = time.time()
        context = "\n".join([chunks[idx] for idx in indices[0] if idx < len(chunks)])
        print(f"Context building time: {time.time() - t3:.2f}s")

        # Add chat history for context-awareness
        if uid not in chat_history:
            chat_history[uid] = []
        history_str = "\n".join([f"User: {q}\nAssistant: {r}" for q, r in chat_history[uid][-5:]])

        t4 = time.time()
        prompt = f"""You are a helpful PDF chatbot. Provide clear, organized answers with bullet points for lists, proper punctuation, and a friendly tone.

**Chat History:**
{history_str}

**Context:**
{context}

**Question:**
{query}"""
        print(f"Prompt building time: {time.time() - t4:.2f}s")

        # Call DeepSeek API
        t5 = time.time()
        headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "You are a helpful PDF chatbot. Provide clear, organized answers with bullet points for lists, proper punctuation, and a friendly tone."},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 400,
            "temperature": 0.7,
            "top_p": 0.9
        }
        response = requests.post("https://api.deepseek.com/v1/chat/completions", json=payload, headers=headers)
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"DeepSeek API error: {response.text}")
        output_text = response.json()["choices"][0]["message"]["content"]
        print(f"DeepSeek API call time: {time.time() - t5:.2f}s")

        t6 = time.time()
        chat_history[uid].append((query, output_text))
        print(f"Post-processing time: {time.time() - t6:.2f}s")

        total_time = time.time() - start_time
        print(f"Total query time: {total_time:.2f}s")
        print(f"User: {uid}")
        print(f"Query: {query}")
        print(f"Output: {output_text}")

        return {"response": output_text}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error querying file: {str(e)}")
    
@app.post("/clear_data")
async def clear_data(request: UIDRequest):
    """
    Delete all files from GCS bucket and Firestore for the given user, keeping the user's folder.
    """
    uid = request.uid
    try:
        # Validate UID
        if not uid:
            raise HTTPException(status_code=400, detail="User ID is required.")

        # Delete files from GCS bucket (contents only)
        bucket = storage_client.bucket(BUCKET_NAME)
        blobs = bucket.list_blobs(prefix=f"{uid}/")
        blob_count = 0
        for blob in blobs:
            blob.delete()
            blob_count += 1
        print(f"Deleted {blob_count} files for user {uid} from GCS.")

        # Delete Firestore data
        user_ref = firestore_client.collection("users").document(uid)
        files_ref = user_ref.collection("files")
        file_count = 0
        chat_count = 0
        for file_doc in files_ref.stream():
            chats_ref = file_doc.reference.collection("chats")
            for chat_doc in chats_ref.stream():
                chat_doc.reference.delete()
                chat_count += 1
            file_doc.reference.delete()
            file_count += 1
        user_ref.delete()
        print(f"Deleted {file_count} files and {chat_count} chats for user {uid} from Firestore.")

        return {"message": "All user data cleared successfully."}
    except Exception as e:
        print(f"Error clearing data for user {uid}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/delete_account")
async def delete_account(request: UIDRequest):
    """
    Delete all user data (GCS, including the user's folder, and Firestore) and the Firebase user account.
    """
    uid = request.uid
    try:
        # Validate UID
        if not uid:
            raise HTTPException(status_code=400, detail="User ID is required.")

        # Delete all GCS data, including the user's folder
        bucket = storage_client.bucket(BUCKET_NAME)
        blobs = bucket.list_blobs(prefix=f"{uid}/")
        blob_count = 0
        for blob in blobs:
            blob.delete()
            blob_count += 1
        print(f"Deleted {blob_count} files and the folder for user {uid} from GCS.")

        # Clear Firestore data
        user_ref = firestore_client.collection("users").document(uid)
        files_ref = user_ref.collection("files")
        file_count = 0
        chat_count = 0
        for file_doc in files_ref.stream():
            chats_ref = file_doc.reference.collection("chats")
            for chat_doc in chats_ref.stream():
                chat_doc.reference.delete()
                chat_count += 1
            file_doc.reference.delete()
            file_count += 1
        user_ref.delete()
        print(f"Deleted {file_count} files and {chat_count} chats for user {uid} from Firestore.")

        # Delete Firebase user
        auth.delete_user(uid)
        print(f"Deleted Firebase user {uid}.")

        return {"message": "Account, folder, and all associated data deleted successfully."}
    except Exception as e:
        print(f"Error deleting account for user {uid}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    import signal
    import sys

    def signal_handler(sig, frame):
        print("Shutting down gracefully...")
        sys.exit(0)

    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)

    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)